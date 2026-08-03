"""
LLM-based document reading for every upload type (jpg/jpeg/png,
pdf, docx, txt, eml/msg) — one Claude call per file handles both
classification and structured extraction, replacing what used to be
two separate systems: Claude for images, and a Tesseract-OCR +
regex-based extractor (dataclasses, hand-written patterns) for
everything else.

That split existed only because the Claude migration originally only
touched the image path. It caused a real inconsistency: the instant
preview shown right after upload could disagree with the official
Databricks-processed result, because the pipeline (see
data-engineering/databricks/2-Text-Extraction.ipynb) already used
Claude for ALL file types, not just images. This file now mirrors
that notebook's approach exactly — same prompt, same per-type field
list, same PDF text-layer-first-then-rendered-pages-as-images
fallback — so the instant preview and the pipeline's real result are
produced the same way and should agree.

Originally ran against a local Ollama server (qwen2.5vl:3b) — fully
private, but the model was heavy enough to make an 8GB M1 MacBook Air
hang under load. Switched to the Claude API (model: Haiku 4.5, the
cheapest Claude model — roughly $0.004 per document at typical image
sizes) to get reliable extraction without taxing local hardware. The
tradeoff: document content now leaves the machine and is sent to
Anthropic's API for processing.
"""

from __future__ import annotations

import base64
import email as email_lib
import json
import logging
import re
import tempfile
from email import policy as email_policy
from io import BytesIO
from typing import Any, Optional

import anthropic
import extract_msg
import fitz  # PyMuPDF — layout-aware PDF text extraction + page rasterization
from charset_normalizer import from_bytes
from docx import Document as DocxDocument
from PIL import Image

logger = logging.getLogger(__name__)

CLAUDE_MODEL = "claude-haiku-4-5-20251001"
MAX_OUTPUT_TOKENS = 2000

# Claude's 10MB limit is checked against the BASE64-ENCODED string,
# not the raw image bytes — confirmed by a real failure where a
# shrunk-to-8.6MB image still got rejected, because base64 inflates
# size by ~4/3 (8.6MB raw -> ~11.5MB encoded, over the 10,485,760
# byte limit). This threshold targets a RAW size whose base64 form
# comfortably clears that: 10,485,760 * 3/4 = 7,864,320 is the exact
# breakeven point, so this stays meaningfully under it. Applied to
# both real uploaded images AND pages rendered from a scanned PDF —
# either one can exceed Claude's limits.
MAX_IMAGE_BYTES = 7_200_000

# Claude also rejects an image if EITHER dimension exceeds 8000px,
# independent of file size — a real failure surfaced this: a highly
# detailed re-encoded JPEG stayed under 10MB but still got a 400
# because its dimensions were too large. Byte-size alone isn't a
# sufficient check.
MAX_IMAGE_DIMENSION = 8000

MEDIA_TYPE_BY_PIL_FORMAT = {
    "JPEG": "image/jpeg",
    "PNG": "image/png",
    "GIF": "image/gif",
    "WEBP": "image/webp",
}

MIN_TEXT_LAYER_LENGTH = 20
# A genuinely scanned/image-only PDF page returns ~0 characters from
# the text layer. Below this length, the page is treated as needing
# the Claude-vision fallback instead of trusting a near-empty extract.


class ImageTooLargeError(Exception):
    """
    Raised when an image (uploaded directly, or rendered from a PDF
    page) can't be shrunk under Claude's limits even at the smallest
    size/quality this function is willing to go to (past that point
    the image would be too degraded to actually read). Given a
    distinct type — rather than a generic Exception — so the caller
    can tell "genuinely too large" apart from other failures and show
    the user a specific, actionable message instead of a generic
    "AI analysis failed".
    """


# ============================================================
# IMAGE HANDLING
# ============================================================


def _shrink_image_if_needed(image_bytes: bytes, media_type: str) -> tuple[bytes, str]:
    """
    Re-encodes as JPEG at progressively lower quality/resolution until
    it's under BOTH Claude's 10MB size limit AND its 8000px dimension
    limit — a file can be small in bytes but still too large in pixels
    (a highly compressed but very high-resolution photo), so both need
    checking, not just size. Only touches the image when it's actually
    over one of the two limits. Raises ImageTooLargeError if it can't
    get small enough without becoming unreadable.
    """

    with Image.open(BytesIO(image_bytes)) as probe:
        original_dimension = max(probe.size)

    if len(image_bytes) <= MAX_IMAGE_BYTES and original_dimension <= MAX_IMAGE_DIMENSION:
        return image_bytes, media_type

    with Image.open(BytesIO(image_bytes)) as image:
        image = image.convert("RGB")

        max_dimension = min(max(image.size), MAX_IMAGE_DIMENSION)
        quality = 90

        while True:
            if max_dimension < max(image.size):
                scale = max_dimension / max(image.size)
                resized = image.resize((max(1, int(image.width * scale)), max(1, int(image.height * scale))))
            else:
                resized = image

            buffer = BytesIO()
            resized.save(buffer, format="JPEG", quality=quality)
            candidate = buffer.getvalue()

            if len(candidate) <= MAX_IMAGE_BYTES:
                return candidate, "image/jpeg"

            # Alternate between dropping quality and shrinking
            # dimensions so a very large, very detailed image doesn't
            # need dozens of iterations to converge.
            if quality > 50:
                quality -= 15
            else:
                max_dimension = int(max_dimension * 0.75)
                quality = 80

            if max_dimension < 400:
                raise ImageTooLargeError(
                    "This image is too large to process, even after compressing it. "
                    "Please upload a smaller file (under 10 MB)."
                )


def _detect_media_type(image_bytes: bytes) -> str:
    """
    Detects the image's real format from its actual bytes rather than
    trusting the filename extension. This matters: a file uploaded as
    "photo.png" can genuinely contain WEBP data (common with browser
    screenshot/export tools) — Claude's API validates the declared
    media_type against the real content and rejects a mismatch, which
    is exactly what happened with a real user upload (filename ended
    in .png, actual bytes were WEBP, Claude returned a 400 error).
    """

    with Image.open(BytesIO(image_bytes)) as image:
        pil_format = image.format

    media_type = MEDIA_TYPE_BY_PIL_FORMAT.get(pil_format or "")
    if media_type is None:
        raise ValueError(f"Unsupported image format for Claude vision: {pil_format}")

    return media_type


# ============================================================
# TEXT EXTRACTION — for born-digital formats (DOCX/TXT/email) and a
# PDF's own embedded text layer. A scanned/image-only PDF page (or an
# image file directly) skips this entirely and goes straight to
# Claude as a picture instead — no OCR anywhere in this file.
# ============================================================


def _extract_text_from_docx(content: bytes) -> str:
    document = DocxDocument(BytesIO(content))

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    tables_text = []
    for table in document.tables:
        for row in table.rows:
            seen_cells = set()
            row_cells = []
            for cell in row.cells:
                if id(cell._tc) in seen_cells:
                    continue
                seen_cells.add(id(cell._tc))
                cell_text = cell.text.strip()
                if cell_text:
                    row_cells.append(cell_text)
            if row_cells:
                tables_text.append(" | ".join(row_cells))

    return "\n".join(paragraphs + tables_text).strip()


def _extract_text_from_txt(content: bytes) -> str:
    best_match = from_bytes(content).best()
    if best_match is None:
        raise ValueError("Could not reliably determine the file's text encoding")
    return str(best_match).strip()


def _strip_html_tags(html: str) -> str:
    text = re.sub(r"<[^>]+>", " ", html)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _parse_eml(content: bytes):
    message = email_lib.message_from_bytes(content, policy=email_policy.default)

    sender = message.get("From")
    recipient = message.get("To")
    subject = message.get("Subject")
    date = message.get("Date")

    body = None
    if message.is_multipart():
        for part in message.walk():
            if part.get_content_type() == "text/plain":
                body = part.get_content()
                break
        else:
            for part in message.walk():
                if part.get_content_type() == "text/html":
                    body = _strip_html_tags(part.get_content())
                    break
    else:
        body = message.get_content()

    return sender, recipient, subject, date, body


def _parse_msg(content: bytes):
    with tempfile.NamedTemporaryFile(suffix=".msg") as tmp_file:
        tmp_file.write(content)
        tmp_file.flush()
        message = extract_msg.Message(tmp_file.name)
        sender = message.sender
        recipient = message.to
        subject = message.subject
        date = str(message.date) if message.date else None
        body = message.body

    return sender, recipient, subject, date, body


def _extract_text_from_email(content: bytes, extension: str) -> str:
    if extension == "eml":
        sender, recipient, subject, date, body = _parse_eml(content)
    else:
        sender, recipient, subject, date, body = _parse_msg(content)

    header_and_body = [
        f"From: {sender or ''}",
        f"To: {recipient or ''}",
        f"Subject: {subject or ''}",
        f"Date: {date or ''}",
        "",
        body or "",
    ]
    return "\n".join(header_and_body).strip()


def _extract_page_text_in_reading_order(page) -> str:
    """
    Reconstructs proper reading order for a page, including
    two-column layouts — PyMuPDF's own block-clustering merges both
    columns' same-row text together otherwise.
    """

    words = page.get_text("words")
    if not words:
        return ""

    lines_by_key: dict[tuple[int, int], dict[str, Any]] = {}
    for x0, y0, x1, y1, word, block_no, line_no, word_no in words:
        key = (block_no, line_no)
        line = lines_by_key.setdefault(key, {"x0": x0, "y0": y0, "x1": x1, "y1": y1, "words": []})
        line["x0"] = min(line["x0"], x0)
        line["y0"] = min(line["y0"], y0)
        line["x1"] = max(line["x1"], x1)
        line["y1"] = max(line["y1"], y1)
        line["words"].append((word_no, word))

    lines = []
    for line in lines_by_key.values():
        text = " ".join(w for _, w in sorted(line["words"]))
        lines.append((line["x0"], line["y0"], line["x1"], line["y1"], text))

    lines.sort(key=lambda l: l[1])

    midpoint = page.rect.width / 2
    left_column, right_column, full_width = [], [], []
    for line in lines:
        x0, x1 = line[0], line[2]
        if x1 <= midpoint:
            left_column.append(line)
        elif x0 >= midpoint:
            right_column.append(line)
        else:
            full_width.append(line)

    if not left_column or not right_column:
        return "\n".join(l[4] for l in lines)

    ordered_lines = []
    left_i = right_i = 0
    for full_width_line in full_width + [None]:
        boundary_y = full_width_line[1] if full_width_line else float("inf")
        while left_i < len(left_column) and left_column[left_i][1] < boundary_y:
            ordered_lines.append(left_column[left_i])
            left_i += 1
        while right_i < len(right_column) and right_column[right_i][1] < boundary_y:
            ordered_lines.append(right_column[right_i])
            right_i += 1
        if full_width_line:
            ordered_lines.append(full_width_line)

    return "\n".join(l[4] for l in ordered_lines)


def _extract_pdf_text_layer(content: bytes) -> tuple[str, bool]:
    """
    Returns (joined_text, needs_vision_fallback). needs_vision_fallback
    is True if ANY page's embedded text layer is too short — a
    document with one real page and one scanned page shouldn't have
    the scanned page silently dropped just because the combined text
    cleared the length threshold.
    """

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        pages_text = []
        needs_fallback = False
        for page in doc:
            page_text = _extract_page_text_in_reading_order(page)
            pages_text.append(page_text)
            if len(page_text) < MIN_TEXT_LAYER_LENGTH:
                needs_fallback = True
        return "\n".join(pages_text).strip(), needs_fallback
    finally:
        doc.close()


def _render_pdf_pages_as_images(content: bytes, dpi: int = 200) -> list[tuple[str, str]]:
    """
    Renders every page of a PDF to a PNG image — used when the PDF's
    text layer is missing/too short (a scanned document), instead of
    running OCR. All pages are sent to Claude together in one call so
    it reads the whole document as a unit. Returns a list of
    (media_type, base64_data) tuples, each already shrunk if needed.
    """

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        images = []
        for page in doc:
            pixmap = page.get_pixmap(dpi=dpi)
            png_bytes = pixmap.tobytes("png")
            png_bytes, media_type = _shrink_image_if_needed(png_bytes, "image/png")
            images.append((media_type, base64.b64encode(png_bytes).decode("utf-8")))
        return images
    finally:
        doc.close()


# ============================================================
# CLASSIFICATION + FIELD EXTRACTION VIA CLAUDE
# ============================================================

EXTRACTION_PROMPT = """You are analyzing a document — either as plain text, as a single image, or as a set of page images from a multi-page scanned document (treat multiple images as one document, read them in order).

Step 1: Decide which ONE of these document types it is: RESUME, EMAIL, INVOICE, BANK_STATEMENT, PRESCRIPTION. If none clearly match, use UNKNOWN.

Step 2: Extract fields based on the type you chose. Read the WHOLE document carefully, including every column if it has a multi-column layout — scan top to bottom in each column separately, do not skip any section.

- RESUME: name, email, phone, linkedin, github, skills (array of strings), education (array of strings), experience (array of strings), projects (array of strings), certifications (array of strings), summary
  - email/phone/linkedin/github are usually near the top or in a contact block — check carefully, they are easy to miss. Look specifically for: a line containing an "@" symbol (email), a line of digits often with dashes (phone), and any line near the contact info that looks like a website handle or URL (linkedin/github), even if it does not start with "http".
  - "skills" = ONLY individual skill/technology names or short skill phrases from a Skills section. NEVER put a job title, company name, or work description in "skills".
  - "experience" = one array entry PER JOB, and each entry must combine the job title, company name, and date range together in one string (e.g. "IT Management Supervisor at Langtown Community College (Sep 2018 - present)"), not just the bare job title alone.
  - "education" = one array entry PER DEGREE/CERTIFICATION found anywhere in an Education or Certifications section, including the institution and year if shown. Include ALL of them, do not stop after the first one or two.
  - "summary" = ONLY the introductory paragraph under a "Summary"/"Objective"/"About" heading. NEVER include contact info, education, or skills text inside "summary".
  - Do not include bullet symbols (•, -, ●, etc.) at the start of any list item — just the text itself.
- EMAIL: sender, recipient, subject, sent_date, body
- INVOICE: invoice_number, invoice_date, due_date, total_amount
- BANK_STATEMENT: account_number, statement_period, opening_balance, closing_balance, transactions (array of strings)
- PRESCRIPTION: patient_name, physician_name, prescription_date, diagnosis, medications (array of strings, include dosage/timing if visible)

Step 3: ALWAYS also include "additional_information" — an array of short strings capturing any OTHER meaningful information visible in the document that isn't already captured by the fields above (for example: a website, a blood group, a reference number, a policy number, a tax ID, a note or remark). Each entry should be a concise "Label: value" string. Do NOT repeat anything already captured in the fields above. Use an empty array if there is genuinely nothing else of note.

Output ONLY a single JSON object, no markdown code fences, no explanation, in exactly this shape:
{"document_type": "<TYPE>", "data": { ...only the fields listed above for that type, plus "additional_information"... }}

Use null for any text field that is not visible in the document, and an empty array for any list field with nothing found.
If document_type is UNKNOWN, set "data" to null.
"""

FIELDS_BY_TYPE = {
    "RESUME": ["name", "email", "phone", "linkedin", "github", "skills", "education", "experience", "projects", "certifications", "summary"],
    "EMAIL": ["sender", "recipient", "subject", "sent_date", "body"],
    "INVOICE": ["invoice_number", "invoice_date", "due_date", "total_amount"],
    "BANK_STATEMENT": ["account_number", "statement_period", "opening_balance", "closing_balance", "transactions"],
    "PRESCRIPTION": ["patient_name", "physician_name", "prescription_date", "diagnosis", "medications"],
}

LIST_FIELDS = {
    "skills", "education", "experience", "projects", "certifications",
    "transactions", "medications", "additional_information",
}


def _parse_llm_json(raw_response: str) -> dict:
    """
    Claude generally respects "no markdown fences", but strip them
    defensively anyway in case a response wraps the JSON in ```json.
    """

    text = raw_response.strip()
    fence_match = re.search(r"```(?:json)?\s*(.*?)\s*```", text, re.DOTALL)
    if fence_match:
        text = fence_match.group(1).strip()

    return json.loads(text)


def _shape_data(document_type: str, raw_data: Optional[dict]) -> Optional[dict]:
    """
    Normalizes the LLM's field values to match the exact shape the
    frontend expects (same fields as the frontend's DocumentData type)
    — fills in missing keys with None/[] rather than trusting the
    model included every key. "additional_information" is appended to
    every type's field list here rather than duplicated into each
    entry in FIELDS_BY_TYPE, since it's the one field common to all 5
    document types.
    """

    if document_type not in FIELDS_BY_TYPE or raw_data is None:
        return None

    expected_fields = [*FIELDS_BY_TYPE[document_type], "additional_information"]
    shaped: dict[str, Any] = {"document_type": document_type}

    for field_name in expected_fields:
        value = raw_data.get(field_name)
        if field_name in LIST_FIELDS:
            shaped[field_name] = value if isinstance(value, list) else []
        else:
            shaped[field_name] = value if isinstance(value, str) and value.strip() else None

    return shaped


def _build_claude_content(content: bytes, extension: str) -> tuple[list[tuple[str, str]], Optional[str]]:
    """
    Returns (images, text) for one file — images is a list of
    (media_type, base64_data) tuples (possibly empty), text is a
    string or None. Exactly one of the two is meaningfully populated
    per file type. Mirrors the Databricks notebook's
    build_claude_request exactly, so the instant preview and the
    pipeline's official result are built from the document the same
    way.
    """

    extension = extension.lower().replace(".", "")

    if extension in ("jpg", "jpeg", "png"):
        media_type = _detect_media_type(content)
        content, media_type = _shrink_image_if_needed(content, media_type)
        return [(media_type, base64.b64encode(content).decode("utf-8"))], None

    elif extension == "pdf":
        text, needs_fallback = _extract_pdf_text_layer(content)
        if needs_fallback:
            return _render_pdf_pages_as_images(content), None
        return [], text

    elif extension == "docx":
        return [], _extract_text_from_docx(content)

    elif extension == "txt":
        return [], _extract_text_from_txt(content)

    elif extension in ("eml", "msg"):
        return [], _extract_text_from_email(content, extension)

    else:
        raise ValueError(f"Unsupported file extension: {extension}")


# A real bank statement upload hit "Expecting ',' delimiter" from
# json.loads — almost certainly an unescaped quote or a literal
# newline inside a transaction-description string (bank statement
# line items are exactly the kind of free text likely to contain
# both). Claude's output isn't deterministic, so one retry with an
# explicit reminder about escaping usually gets a clean parse the
# second time, without needing to guess-repair the malformed JSON
# text itself.
JSON_RETRY_REMINDER = (
    "\n\nIMPORTANT: Your previous response could not be parsed as valid JSON. Output ONLY a "
    "single, strictly valid JSON object — escape every double quote and backslash inside "
    'string values (\\" and \\\\), and never include a literal newline inside a string value '
    "(use \\n instead). No text before or after the JSON."
)


def extract_with_llm(content: bytes, filename: str, api_key: str) -> dict:
    """
    Sends the file to the Claude API (Haiku 4.5) and returns a dict
    shaped like {"document_type": ..., "classification_confidence":
    ..., "data": ..., "raw_text": ""} — matching QuickResult in
    upload.py. Handles every supported extension (image, PDF, DOCX,
    TXT, EML/MSG) — see _build_claude_content. Raises on any failure
    (bad/missing API key, network error, malformed response, etc.) —
    the caller (upload.py) is responsible for catching and returning
    an explicit error rather than silently degrading.
    """

    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    images, text = _build_claude_content(content, extension)

    image_blocks: list[dict] = [
        {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": b64_data}}
        for media_type, b64_data in images
    ]

    prompt = EXTRACTION_PROMPT
    if text:
        prompt = prompt + "\n\nDocument text:\n" + text

    client = anthropic.Anthropic(api_key=api_key)

    raw_text = ""
    parsed: Optional[dict] = None
    last_error: Optional[json.JSONDecodeError] = None

    for attempt, attempt_prompt in enumerate((prompt, prompt + JSON_RETRY_REMINDER)):
        response = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=MAX_OUTPUT_TOKENS,
            messages=[{"role": "user", "content": [*image_blocks, {"type": "text", "text": attempt_prompt}]}],
        )
        raw_text = "".join(block.text for block in response.content if block.type == "text")

        try:
            parsed = _parse_llm_json(raw_text)
            break
        except json.JSONDecodeError as exc:
            last_error = exc
            logger.warning(
                "Claude returned malformed JSON on attempt %d for %s: %s | raw response (truncated): %s",
                attempt + 1, filename, exc, raw_text[:2000],
            )

    if parsed is None:
        # Logged with the full raw text above (twice, once per failed
        # attempt) so a future occurrence is actually diagnosable —
        # previously nothing captured what Claude had actually sent,
        # only the terse json.loads error message.
        raise last_error

    document_type = parsed.get("document_type", "UNKNOWN")
    if document_type not in FIELDS_BY_TYPE:
        document_type = "UNKNOWN"

    data = _shape_data(document_type, parsed.get("data"))

    return {
        "document_type": document_type,
        "classification_confidence": 1.0 if document_type != "UNKNOWN" else 0.0,
        "data": data,
        "raw_text": "",
    }


def process_document_fast(content: bytes, filename: str, anthropic_api_key: Optional[str] = None) -> dict:
    """
    The single entry point the upload endpoint calls — runs
    classification + field extraction on one uploaded file
    synchronously (via Claude), for every supported file type. Never
    raises: any failure just means no instant preview for this
    upload, which is not fatal (the Databricks pipeline still runs
    and produces the real result). Errors are returned, not thrown,
    so the upload endpoint can decide how to surface them.
    """

    if not anthropic_api_key:
        return {"error": "AI analysis unavailable: no Claude API key configured"}

    try:
        return extract_with_llm(content, filename, anthropic_api_key)
    except ImageTooLargeError as exc:
        # Not a real failure worth a stack trace — an expected,
        # user-actionable case (their file was too big).
        return {"error": str(exc)}
    except Exception as exc:
        logger.exception("LLM extraction failed for upload: %s", filename)
        return {"error": f"AI analysis failed: {exc}"}
